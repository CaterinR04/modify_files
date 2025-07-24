# Requerimientos: Python3, vscode (editor de codigo), pip install python-docx
# Ejecutar: python3 script.py #
import os
import re
from docx import Document
from docx.shared import RGBColor

PATTERN_EXP_CONF = re.compile(r'\bExponencial\s+Confirming\b(?!\s+S\.A\.S\b)', re.IGNORECASE)
PATTERN_EXP      = re.compile(r'\bExponencial\b(?!\s+Confirming\s+S\.A\.S\b)', re.IGNORECASE)

def process_docx(path):
    doc = Document(path)
    def format_and_replace_para(para):
        text = para.text
        text = PATTERN_EXP_CONF.sub("Mente", text)
        text = PATTERN_EXP.sub("Mente", text)
        for run in para.runs:
            run.text = ""
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = "Museo Sans 500"
    for para in doc.paragraphs:
        format_and_replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    format_and_replace_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            format_and_replace_para(para)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        format_and_replace_para(para)
        for para in section.footer.paragraphs:
            format_and_replace_para(para)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        format_and_replace_para(para)
    dir_name, filename = os.path.split(path)
    base, ext = os.path.splitext(filename)
    new_filename = f"{base}_modificado{ext}"
    new_path = os.path.join(dir_name, new_filename)
    doc.save(new_path)
    print(f"Guardado: {new_filename} (original: {filename})")

def main(folder_path):
    for fn in os.listdir(folder_path):
        if fn.lower().endswith(".docx"):
            process_docx(os.path.join(folder_path, fn))

if __name__ == "__main__":
    carpeta = r"./files"
    main(carpeta)
